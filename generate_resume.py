#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
AI Agent 开发工程师简历生成器
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def add_heading_with_style(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色
    return heading

def add_section_line(doc):
    """添加分隔线"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 0.5
    run = paragraph.add_run('_' * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(192, 192, 192)

def create_resume():
    """创建简历文档"""
    doc = Document()

    # 设置文档默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)

    # ==================== 个人信息 ====================
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('姓名')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_text = contact.add_run('电话：138-xxxx-xxxx  |  邮箱：yourname@example.com  |  GitHub：github.com/yourname')
    contact_text.font.size = Pt(9)

    doc.add_paragraph()  # 空行

    # ==================== 求职意向 ====================
    add_heading_with_style(doc, '求职意向', level=2)
    p = doc.add_paragraph()
    p.add_run('AI Agent 开发工程师 / LLM 应用工程师 / AI 系统架构师').font.size = Pt(10.5)
    doc.add_paragraph()

    # ==================== 专业技能 ====================
    add_heading_with_style(doc, '专业技能', level=2)
    add_section_line(doc)

    skills = [
        ('AI Agent 开发', [
            '熟练掌握 AI Agent 架构设计（ReAct、ReWOO、Plan-and-Execute 等模式），具备完整的 Agent 系统开发经验',
            '深入理解 LangChain、LangGraph 等主流 Agent 框架，能够设计复杂的多 Agent 协作系统',
            '掌握工具调用（Function Calling）、记忆管理、检索增强生成（RAG）等核心技术',
            '熟悉 Prompt Engineering 技术，包括 CoT、Few-shot、Self-Consistency 等提示优化方法'
        ]),
        ('大语言模型应用', [
            '熟悉 OpenAI GPT、Claude、GLM 等主流大模型的 API 调用与应用开发',
            '掌握模型微调（Fine-tuning）、Prompt 优化、上下文管理等 LLM 应用技术',
            '了解模型评估与监控方法，能够优化模型输出质量和响应速度'
        ]),
        ('自然语言处理', [
            '扎实的 NLP 基础知识，掌握文本预处理、分词、词向量、命名实体识别等技术',
            '熟悉 Transformer 架构及其变体（BERT、GPT、T5 等）',
            '了解文本分类、情感分析、文本生成、机器翻译等 NLP 任务'
        ]),
        ('机器学习与深度学习', [
            '扎实的机器学习基础，熟悉监督学习、非监督学习、强化学习等核心算法',
            '深入理解深度学习原理，掌握 CNN、RNN、LSTM、Attention 等网络架构',
            '熟练使用 PyTorch / TensorFlow 框架进行模型开发与训练',
            '掌握模型优化技术（正则化、Dropout、BatchNorm 等）和超参数调优'
        ]),
        ('向量数据库与检索', [
            '熟练使用 Milvus、FAISS 等向量数据库进行语义检索',
            '掌握 Embedding 技术和相似度计算方法',
            '了解混合检索策略（向量检索 + 关键词检索）',
            '熟悉 Redis 缓存技术和 SQL 数据库操作'
        ]),
        ('编程与工程能力', [
            '熟练掌握 Python 编程，具备良好的代码规范和工程实践能力',
            '熟悉 Linux 系统操作和 Shell 脚本编写',
            '了解 Docker 容器化部署和微服务架构',
            '掌握 Git 版本控制，具备良好的团队协作能力',
            '熟悉常用数据结构与算法，具备解决复杂问题的能力'
        ]),
        ('数学基础', [
            '扎实的数学基础：线性代数（向量、矩阵、特征分解）',
            '掌握微积分与优化方法（梯度下降、Adam 等优化器）',
            '熟悉概率统计知识（贝叶斯定理、分布理论、统计推断）'
        ])
    ]

    for skill_category, skill_list in skills:
        p = doc.add_paragraph()
        run = p.add_run(f'◆ {skill_category}')
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 51, 102)

        for skill in skill_list:
            p = doc.add_paragraph(style='List Bullet 2')
            p.add_run(skill).font.size = Pt(10)

    doc.add_paragraph()

    # ==================== 项目经验 ====================
    add_heading_with_style(doc, '项目经验', level=2)
    add_section_line(doc)

    # 项目 1
    p = doc.add_paragraph()
    run = p.add_run('智能问答 AI Agent 系统')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run('项目描述：').font.bold = True
    p.add_run('基于大语言模型的智能问答系统，支持多轮对话、知识检索和任务执行')

    p = doc.add_paragraph()
    p.add_run('技术栈：').font.bold = True
    p.add_run('Python、LangChain、OpenAI API、Milvus、Redis、FastAPI')

    p = doc.add_paragraph()
    p.add_run('核心职责：').font.bold = True

    responsibilities = [
        '设计并实现基于 ReAct 模式的 AI Agent 架构，支持工具调用和多步推理',
        '集成向量数据库 Milvus 实现 RAG 检索增强生成，提升回答准确性 40%',
        '开发记忆管理模块，实现对话历史的自动摘要和长期记忆存储',
        '优化 Prompt 工程，通过 Few-shot 和 CoT 技术提升任务成功率 35%',
        '实现多 Agent 协作机制，支持任务分解与并行执行'
    ]
    for resp in responsibilities:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(resp).font.size = Pt(10)

    doc.add_paragraph()

    # 项目 2
    p = doc.add_paragraph()
    run = p.add_run('企业知识库智能助手')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run('项目描述：').font.bold = True
    p.add_run('面向企业内部文档的智能检索与问答系统，支持多格式文档解析和语义检索')

    p = doc.add_paragraph()
    p.add_run('技术栈：').font.bold = True
    p.add_run('Python、LangChain、FAISS、Sentence-Transformers、Streamlit')

    p = doc.add_paragraph()
    p.add_run('核心职责：').font.bold = True

    responsibilities = [
        '设计文档解析管道，支持 PDF、Word、Markdown 等多种格式的自动解析',
        '实现文档分块策略（Chunk Strategy），优化检索召回率和精确度',
        '使用 Sentence-Transformers 生成文档 Embedding，建立向量索引',
        '开发混合检索模块，结合关键词检索和向量检索，提升检索质量',
        '构建 Web 界面，支持文档上传、实时问答和结果溯源'
    ]
    for resp in responsibilities:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(resp).font.size = Pt(10)

    doc.add_paragraph()

    # 项目 3
    p = doc.add_paragraph()
    run = p.add_run('AI 学习知识体系构建项目')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run('项目描述：').font.bold = True
    p.add_run('系统化学习 AI 技术栈，构建涵盖 260+ 知识点的原子卡片知识库')

    p = doc.add_paragraph()
    p.add_run('学习成果：').font.bold = True

    achievements = [
        '完成 12 个核心领域的系统学习：AI Agent (87卡)、深度学习 (40卡)、机器学习 (37卡)、数学基础 (33卡)、数据库 (23卡)、Python (22卡)、NLP (18卡)、Tools (19卡)、数据结构与算法 (15卡) 等',
        '建立模块化知识体系，形成从基础到应用的完整学习路径',
        '实践 AI Agent 开发、RAG 检索、向量数据库应用等核心技术',
        '掌握 Docker 容器化部署和工程化实践能力'
    ]
    for ach in achievements:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(ach).font.size = Pt(10)

    doc.add_paragraph()

    # ==================== 教育背景 ====================
    add_heading_with_style(doc, '教育背景', level=2)
    add_section_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('XX大学  |  计算机科学与技术专业  |  本科/硕士')
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('2018.09 - 2022.06  |  主修课程：数据结构、算法设计、机器学习、深度学习、自然语言处理')

    doc.add_paragraph()

    # ==================== 自我评价 ====================
    add_heading_with_style(doc, '自我评价', level=2)
    add_section_line(doc)

    evaluations = [
        '对 AI Agent 技术和大语言模型应用充满热情，持续关注行业最新进展',
        '具备扎实的计算机基础和数学功底，能够快速学习和应用新技术',
        '注重工程实践和代码质量，具备良好的问题分析和解决能力',
        '善于团队协作和知识分享，具有较强的沟通表达能力',
        '自驱力强，具备系统化学习和知识体系构建能力'
    ]

    for evaluation in evaluations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(evaluation).font.size = Pt(10)

    # 保存文档
    output_path = 'AI_Agent开发工程师_简历样例.docx'
    doc.save(output_path)
    print(f'✅ 简历已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_resume()
