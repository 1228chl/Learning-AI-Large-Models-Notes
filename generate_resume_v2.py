#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
AI Agent 开发工程师简历生成器 V2 - 基于专业模板建议
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def add_heading_with_style(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_section_line(doc):
    """添加分隔线"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 0.5
    run = paragraph.add_run('_' * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(192, 192, 192)

def create_resume():
    """创建简历文档"""
    doc = Document()

    # 设置文档默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)

    # ==================== 个人信息 ====================
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('[你的姓名]')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_text = contact.add_run('手机号码：138-xxxx-xxxx  |  电子邮箱：yourname@example.com')
    contact_text.font.size = Pt(9)

    contact2 = doc.add_paragraph()
    contact2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact2_text = contact2.add_run('GitHub：github.com/yourname  |  个人作品：[Agent项目展示链接]')
    contact2_text.font.size = Pt(9)
    contact2_text.font.color.rgb = RGBColor(0, 102, 204)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('求职意向：AI Agent 开发工程师 / 大模型应用开发工程师')
    run.font.size = Pt(10.5)
    run.font.bold = True

    doc.add_paragraph()

    # ==================== 专业技能 ====================
    add_heading_with_style(doc, '专业技能', level=2)
    add_section_line(doc)

    p = doc.add_paragraph()
    p.add_run('请根据实际掌握程度调整顺序，将最擅长的、与 Agent 最相关的排在前面').font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph()

    skills = [
        ('智能体框架', '精通 LangChain/LangGraph、LlamaIndex，熟练掌握 AutoGen、CrewAI 等多智能体协作框架；深入理解 ReAct、Plan-and-Execute、ReWOO 等 Agent 推理范式。'),

        ('大模型应用', '熟悉 GPT-4o、Claude 3.5 Sonnet、DeepSeek-V3 等主流模型 API 调用与原理；掌握提示工程（CoT、Few-shot、Self-Consistency、结构化输出）；有 Function Calling / Tool Use 实战经验。'),

        ('检索增强生成', '扎实的 RAG 全链路实战能力，包括文档解析（PyMuPDF/Unstructured）、切片策略（Semantic/Recursive）、向量嵌入（BGE/Text-Embedding-3）、向量数据库（Milvus/Chroma/FAISS）及检索优化（重排序、混合检索、Self-RAG）。'),

        ('工具集成与环境交互', '熟练实现 Agent 工具定制，包括 API 调用、代码解释器（Sandbox）、SQL 自动生成与执行；了解 MCP 协议以构建标准化的工具生态；能设计高可用工具调用机制（异常容错、重试策略）。'),

        ('模型工程化与服务', '精通 Python，擅长使用 FastAPI/Flask 封装 Agent 服务；熟悉 Docker 容器化部署；能够利用 LangSmith/LangFuse 进行 Agent 全链路追踪与效果评估；了解流式输出（SSE）和并发流控。'),

        ('记忆与状态管理', '掌握对话记忆（Buffer/Summary）与长期记忆的实现（基于数据库存储用户画像/知识库），懂得设计多轮对话中的上下文压缩策略；熟悉 Redis 缓存技术。'),

        ('深度学习与 NLP 基础', '扎实的深度学习基础，熟悉 Transformer 架构及其变体（BERT、GPT、T5）；掌握注意力机制、位置编码等核心技术；了解模型训练与优化方法（梯度下降、Adam、正则化）。'),

        ('数学与算法基础', '扎实的数学基础（线性代数、微积分、概率统计）；熟悉常用数据结构与算法；具备解决复杂问题的能力。'),

        ('模型微调与优化（加分项）', '了解 LoRA/QLoRA 微调技术，有使用 Hugging Face Transformers、LLaMA-Factory 对开源模型进行指令微调的经验。')
    ]

    for skill_category, skill_detail in skills:
        p = doc.add_paragraph()
        run = p.add_run(f'• {skill_category}：')
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 51, 102)

        run2 = p.add_run(skill_detail)
        run2.font.size = Pt(10)

    doc.add_paragraph()

    # ==================== 项目经历 ====================
    add_heading_with_style(doc, '项目经历', level=2)
    add_section_line(doc)

    p = doc.add_paragraph()
    p.add_run('这是简历的重头戏！务必突出"智能体"决策能力和自动化成果，多用量化指标。').font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph()

    # 项目 1
    p = doc.add_paragraph()
    run = p.add_run('项目一：私人 AI 旅行规划师 Agent')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run('项目描述：').font.bold = True
    p.add_run('针对"想做攻略怕麻烦"的痛点，开发了一个可进行多轮对话、调用多平台实时数据的旅行规划智能体。')

    p = doc.add_paragraph()
    p.add_run('技术栈：').font.bold = True
    p.add_run('LangGraph、GPT-4o、SerpAPI、Google Maps API、Streamlit')

    p = doc.add_paragraph()
    p.add_run('职责与成果：').font.bold = True

    items = [
        '设计 Plan-and-Execute 架构，将用户模糊需求分解为查天气、搜景点、算路径、排行程四个原子任务',
        '自定义 5 个高可用工具（天气查询、机票比价等），处理 API 异常容错，使 Agent 工具调用成功率达到 99%',
        '设计带总结功能的记忆模块，支持用户通过"不爬山"、"换成川菜"等指令进行多轮行程微调，用户满意度评分 4.6/5.0'
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()

    # 项目 2
    p = doc.add_paragraph()
    run = p.add_run('项目二：私人知识库深度问答 Agent（RAG 深化）')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run('项目描述：').font.bold = True
    p.add_run('实现对私有 PDF/Word 文档库的"律师级"精确问答，需定位原文，杜绝幻觉。')

    p = doc.add_paragraph()
    p.add_run('技术栈：').font.bold = True
    p.add_run('LlamaIndex、BGE-M3、Milvus、BGE Reranker、FastAPI')

    p = doc.add_paragraph()
    p.add_run('职责与成果：').font.bold = True

    items = [
        '构建非结构化文档处理管道，实现父子文档切分策略以保留上下文语义',
        '采用 Self-RAG 机制，让 Agent 在回答问题前先通过自我反思判断是否需要检索，在证据不足时明确回答"不知道"，将幻觉率控制在 2% 以下',
        '使用 FastAPI 将整套逻辑封装为 API 并上云部署，支持 200 用户同时在线'
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()

    # 项目 3
    p = doc.add_paragraph()
    run = p.add_run('项目三：多智能体协作研究系统')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run('项目描述：').font.bold = True
    p.add_run('针对复杂研究任务，设计"计划制定-数据采集-深度分析-报告生成"多 Agent 流水线。')

    p = doc.add_paragraph()
    p.add_run('技术栈：').font.bold = True
    p.add_run('AutoGen / LangGraph、GPT-4o、Redis、Docker')

    p = doc.add_paragraph()
    p.add_run('职责与成果：').font.bold = True

    items = [
        '基于 AutoGen/LangGraph 设计多智能体协作架构，实现 Planner、Researcher、Analyst、Reporter 四个专业 Agent 的协同工作',
        '使用 ReAct 模式让 Agent 自主判断调用工具（搜索引擎、数据库、代码执行器）还是进行推理',
        '处理长流程任务的时间从人工的 3 小时缩短到 15 分钟，最终报告质量达到人类专家水平的 85%',
        '通过 LangSmith 追踪 Agent 思考链路，优化提示词和工具调用策略，将任务成功率从 72% 提升至 94%'
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()

    # 项目 4
    p = doc.add_paragraph()
    run = p.add_run('项目四：AI 学习知识体系构建（260+ 原子卡片）')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run('项目描述：').font.bold = True
    p.add_run('系统化学习 AI 技术栈，构建涵盖 260+ 知识点的模块化知识库，覆盖从数学基础到 Agent 应用的完整路径。')

    p = doc.add_paragraph()
    p.add_run('学习成果：').font.bold = True

    items = [
        'AI Agent 领域（87 张卡片）：深入学习 ReAct、Plan-and-Execute、Multi-Agent 等架构模式，掌握 LangChain/LangGraph 框架实战',
        '深度学习（40 张）+ 机器学习（37 张）：系统掌握神经网络、Transformer、注意力机制等核心技术',
        'NLP 技术（18 张）：熟悉文本处理、词向量、预训练模型、文本生成等技术栈',
        '数据库与检索（23 张）：掌握 Milvus 向量数据库、Redis 缓存、SQL 数据库及混合检索策略',
        '数学基础（33 张）：扎实的线性代数、微积分、概率统计基础，支撑深度学习理论理解',
        '工程实践（19 张）：熟悉 Docker 容器化、Python 工程化、网络部署等实战技能'
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()

    # ==================== 工作/实习经历（可选） ====================
    add_heading_with_style(doc, '工作经历 / 实习经历（可选）', level=2)
    add_section_line(doc)

    p = doc.add_paragraph()
    p.add_run('如有相关工作或实习经历，请在此填写。务必突出"智能体"决策能力和自动化成果。').font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('[公司名称] | AI Agent 开发工程师 / 实习生 | 202X.XX - 202X.XX')
    run.font.bold = True
    run.font.size = Pt(11)

    items = [
        '多智能体协作系统搭建：针对[具体业务场景]，基于 AutoGen/LangGraph 设计多 Agent 流水线，处理时间从 X 小时缩短到 X 分钟，效率提升 XX%',
        '智能客服 Agent 重构：使用 LangChain + ReAct 模式，让 Agent 自主判断调用工具，独立解决率提升至 92%',
        '高可用 Agent 网关服务：基于 FastAPI 统一管理模型路由、Tool 鉴权与并发流控，支持 500 QPS 并发调用，P99 延迟低于 2s',
        '评估与迭代机制建设：引入 LangSmith 追踪 Agent 链路，构建自动化回归测试集，将线上准确率从 78% 提升至 94%'
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet 2')
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()

    # ==================== 教育背景 ====================
    add_heading_with_style(doc, '教育背景', level=2)
    add_section_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('[大学名称] | [专业] | [本科/硕士] | 202X.XX - 202X.XX')
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('主修课程：数据结构与算法、机器学习、深度学习、自然语言处理、数据库系统')

    p = doc.add_paragraph()
    p.add_run('如果绩点高或拿过奖学金可以标注；跨专业建议写上辅修或自学的相关课程').font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ==================== 证书与开源贡献（可选） ====================
    add_heading_with_style(doc, '证书与开源贡献（可选）', level=2)
    add_section_line(doc)

    p = doc.add_paragraph()
    p.add_run('开源贡献：').font.bold = True
    p.add_run('为 LangChain 仓库提交 PR 并被 Merge；自研开源项目"[项目名称]"获得 XXX+ Star')

    p = doc.add_paragraph()
    p.add_run('证书：').font.bold = True
    p.add_run('DeepLearning.AI 的《AI Agentic Design Patterns with AutoGen》、吴恩达《LangChain for LLM Application Development》等')

    doc.add_paragraph()

    # ==================== 专家建议（红字提醒） ====================
    add_heading_with_style(doc, '📌 专家评审建议（填写时请删除此部分）', level=2)
    add_section_line(doc)

    tips = [
        '一定要体现"自主决策"逻辑：普通 AI 工程师写"调用 GPT API"，Agent 工程师要写"让模型自主判断选择工具，并在失败时自我纠错重试"',
        '工程化能力是区分度：提到 Docker、FastAPI 封装、流式输出（SSE）、断线重连、LangSmith 监控等',
        '"避坑"经验很值钱：如何解决 Agent 死循环？如何降低 RAG 幻觉？如何处理 Token 超限？',
        '量化指标必不可少：提升 XX%、降低 XX%、支持 XXX QPS、准确率达到 XX%',
        '关键词布局：ReAct、CoT、多智能体、记忆管理、工具调用、Self-RAG、Plan-and-Execute'
    ]

    for tip in tips:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(tip)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 0, 0)

    # 保存文档
    output_path = 'AI_Agent开发工程师_简历样例_V2.docx'
    doc.save(output_path)
    print(f'✅ 简历已生成：{output_path}')
    print(f'📝 基于专业模板建议，更符合行业标准')
    print(f'⚠️  请仔细阅读红色提示，根据实际情况填写')
    return output_path

if __name__ == '__main__':
    create_resume()
